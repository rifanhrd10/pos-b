from pathlib import Path
from copy import deepcopy

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/rifan/APPS SYSTEM/pos-b")
OUT = ROOT / "Bayaro_POS_Manual_Book.docx"

NAVY = "0A1F54"
BLUE = "135FEF"
AQUA = "22D3EE"
INK = "172033"
MUTED = "5C6B82"
PALE = "F4F8FF"
PALE_BLUE = "E8F0FF"
GREEN = "E7F8F2"
ORANGE = "FFF4E5"
WHITE = "FFFFFF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Bayaro POS  |  Manual Book  •  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def style_run(run, size=None, bold=None, color=INK, italic=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, space_after=6, color=INK, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    style_run(r, size=size, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    style_run(r, size=10, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    style_run(r, size=10, color=INK)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r, size=20, bold=True, color=NAVY)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r, size=13, bold=True, color=BLUE)
    return p


def add_callout(doc, title, text, color=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_col_width(cell, 6.35)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    shade(cell, color)
    set_cell_border(cell, top={"val": "single", "sz": 6, "color": "BFD2F6"}, bottom={"val": "single", "sz": 6, "color": "BFD2F6"}, left={"val": "single", "sz": 6, "color": "BFD2F6"}, right={"val": "single", "sz": 6, "color": "BFD2F6"})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    style_run(r, size=10, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    style_run(r2, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_col_width(cell, widths[i])
        set_cell_margins(cell)
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        style_run(r, size=9, bold=True, color=WHITE)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_col_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, "F8FAFE" if len(table.rows) % 2 == 0 else WHITE)
            set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": "D9E2F2"})
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            style_run(r, size=8.7, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_screenshot(doc, filename, caption, width=6.35, crop_height=930):
    path = ROOT / filename
    if not path.exists():
        return
    image_path = path
    temp_path = ROOT / (".manual-crop-" + path.name)
    with Image.open(path) as im:
        if im.height > crop_height:
            crop = im.crop((0, 0, im.width, crop_height))
            crop.save(temp_path)
            image_path = temp_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    style_run(r, size=8.5, italic=True, color=MUTED)
    if temp_path.exists():
        temp_path.unlink()


def add_section_cover(doc, number, title, summary):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"BAGIAN {number}")
    style_run(r, size=9, bold=True, color=AQUA)
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(7)
    h.paragraph_format.keep_with_next = True
    rh = h.add_run(title)
    style_run(rh, size=24, bold=True, color=NAVY)
    add_body(doc, summary, space_after=12, color=MUTED, size=11)


def page_break(doc):
    doc.add_page_break()


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    for name, size, color in (("Heading 1", 20, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11, NAVY)):
        s = doc.styles[name]
        s.font.name = "Arial"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
    footer = sec.footer.paragraphs[0]
    add_page_number(footer)


def build():
    doc = Document()
    configure(doc)

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(30)
    logo = ROOT / "public/branding/bayaro-logo-transparent.png"
    if logo.exists():
        cover.add_run().add_picture(str(logo), width=Inches(2.25))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run("MANUAL BOOK")
    style_run(r, size=12, bold=True, color=AQUA)
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run("Bayaro POS")
    style_run(r, size=34, bold=True, color=NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    r = sub.add_run("Panduan penggunaan aplikasi kasir dan administrasi bisnis")
    style_run(r, size=13, color=MUTED)
    add_callout(doc, "Versi dokumen", "Disusun berdasarkan aplikasi Bayaro POS yang berjalan di localhost, data demo, dan hak akses yang tersedia di konfigurasi aplikasi.", color=PALE)
    add_body(doc, "Dokumen ini membantu Owner, Admin, Manager, Kasir, Staff Gudang, dan IT Admin memahami alur kerja Bayaro POS dari login sampai pelaporan.", space_after=6, size=11)
    add_body(doc, "Tanggal penyusunan: 13 Juli 2026", space_after=0, color=MUTED, size=9.5)
    page_break(doc)

    # Contents and quick start
    add_h1(doc, "Daftar Isi")
    toc = [
        "1. Gambaran umum dan persiapan awal",
        "2. Login dan jalur akses kasir",
        "3. Hak akses dan role",
        "4. Dashboard",
        "5. POS Kasir dan transaksi",
        "6. Produk dan kategori",
        "7. Inventori dan stok",
        "8. Pelanggan dan promo",
        "9. Laporan",
        "10. Operasional: karyawan, outlet, dan role",
        "11. Pengaturan dan akun",
        "12. Alur kerja harian dan troubleshooting singkat",
    ]
    for item in toc:
        add_bullet(doc, item)
    add_h2(doc, "Informasi aplikasi")
    add_table(doc, ["Item", "Nilai"], [
        ["Alamat lokal", "http://localhost:3000"],
        ["Nama bisnis demo", "Bayaro Coffee Demo"],
        ["Kode toko demo", "BAYARO01"],
        ["Paket demo", "Pro (TRIAL)"],
        ["Database", "PostgreSQL melalui Prisma"],
    ], [1.65, 4.7])
    add_callout(doc, "Catatan keamanan", "Akun demo di manual ini hanya untuk lingkungan lokal/demo. Ganti password, PIN, AUTH_SECRET, dan kredensial gateway sebelum aplikasi digunakan di lingkungan produksi.", ORANGE)

    # Login and access
    page_break(doc)
    add_section_cover(doc, "1", "Login dan jalur akses", "Bayaro POS menyediakan login admin berbasis email/password dan jalur kasir publik berbasis kode toko, outlet, dan PIN.")
    add_screenshot(doc, "manual-assets-login.png", "Gambar 1. Halaman login Bayaro POS.")
    add_h2(doc, "Login Owner, Admin, atau Manager")
    add_number(doc, "Buka http://localhost:3000/login.")
    add_number(doc, "Masukkan email dan password sesuai akun yang diberikan.")
    add_number(doc, "Klik Masuk ke Dashboard.")
    add_number(doc, "Jika akun memiliki data bisnis, aplikasi membuka Dashboard. Jika belum ada bisnis, aplikasi mengarahkan ke onboarding.")
    add_h2(doc, "Akun demo")
    add_table(doc, ["Role", "Email", "Password", "PIN kasir"], [
        ["Admin / Owner", "admin@bayaro.id", "demo123", "1234"],
        ["Manager", "manager@bayaro.id", "demo123", "2222"],
        ["Kasir", "kasir@bayaro.id", "demo123", "3333"],
        ["IT Admin", "itadmin@bayaro.id", "admin123", "-"],
    ], [1.3, 2.2, 1.35, 1.5])
    add_h2(doc, "Jalur kasir publik")
    add_body(doc, "Jalur ini cocok untuk kasir yang tidak perlu masuk ke dashboard admin. Kasir hanya membutuhkan kode toko, outlet, dan PIN 4 digit.")
    add_screenshot(doc, "manual-assets-kasir-enter.png", "Gambar 2. Memulai jalur kasir publik dengan kode toko.")
    add_number(doc, "Buka /kasir/enter atau klik Masuk sebagai Kasir dari halaman login.")
    add_number(doc, "Masukkan kode toko, contoh BAYARO01, lalu klik Lanjutkan.")
    add_screenshot(doc, "manual-assets-kasir-outlets.png", "Gambar 3. Memilih outlet tempat kasir bertugas.")
    add_number(doc, "Pilih outlet yang sesuai.")
    add_screenshot(doc, "manual-assets-kasir-pin.png", "Gambar 4. Memasukkan PIN kasir 4 digit.")
    add_number(doc, "Masukkan PIN 4 digit. Setelah berhasil, aplikasi membuka POS Kasir.")

    # Roles
    page_break(doc)
    add_section_cover(doc, "2", "Hak akses dan role", "Hak akses mengatur menu yang terlihat dan tindakan yang boleh dilakukan oleh setiap pengguna.")
    add_table(doc, ["Role", "Fokus akses", "Kemampuan utama"], [
        ["Owner", "Akses penuh", "Dashboard, POS, CRUD produk, stok, laporan, pelanggan, promo, karyawan, outlet, role, dan seluruh pengaturan."],
        ["Admin", "Administrasi bisnis", "Semua fitur bisnis kecuali mengelola permission role."],
        ["Manager", "Operasional cabang", "Dashboard, POS, diskon manual, lihat produk/stok/laporan, pelanggan, karyawan, dan outlet."],
        ["Kasir", "Penjualan", "Akses POS, lihat produk, dan cari pelanggan."],
        ["Warehouse", "Inventori", "Lihat produk/stok dan melakukan penyesuaian, transfer, atau stok opname."],
        ["IT Admin", "Sistem platform", "Mengelola bisnis/toko, paket, subscription, monitoring, dan system. Tidak memakai menu bisnis biasa."],
    ], [1.0, 1.55, 3.8])
    add_h2(doc, "Permission penting")
    add_table(doc, ["Kelompok", "Permission", "Dampak"], [
        ["Dashboard", "dashboard.view", "Melihat ringkasan KPI dan aktivitas bisnis."],
        ["POS", "pos.access, pos.discount, pos.void, pos.refund, pos.close_shift", "Mengakses kasir, diskon, pembatalan/refund, dan tutup shift."],
        ["Produk", "products.view, products.manage, products.pricing", "Melihat produk, CRUD produk, dan mengubah harga."],
        ["Inventori", "inventory.view, inventory.manage", "Melihat stok dan melakukan mutasi stok."],
        ["Laporan", "reports.view, reports.export", "Melihat dan mengekspor laporan."],
        ["Master data", "customers.*, promos.*, employees.*", "Mengelola pelanggan, promo, dan karyawan sesuai hak akses."],
        ["Settings", "settings.manage, settings.roles", "Mengelola konfigurasi bisnis dan role/permission."],
    ], [1.25, 2.6, 2.5])

    # Dashboard
    add_section_cover(doc, "3", "Dashboard", "Dashboard adalah halaman ringkasan untuk memantau omzet, transaksi, produk, pembayaran, dan stok rendah.")
    add_screenshot(doc, "manual-assets-dashboard.png", "Gambar 5. Dashboard dengan KPI, grafik omzet, transaksi terbaru, dan status stok.")
    add_h2(doc, "Bagian utama")
    add_bullet(doc, "KPI: Pendapatan, Transaksi, Rata-rata Order, dan Item Terjual.")
    add_bullet(doc, "Pemilih periode: Hari ini, 7 Hari, atau 30 Hari.")
    add_bullet(doc, "Grafik omzet: melihat pola pendapatan pada periode terpilih.")
    add_bullet(doc, "Produk terlaris dan metode pembayaran: analisis komposisi penjualan.")
    add_bullet(doc, "Transaksi terbaru: membuka daftar penjualan lebih lengkap melalui Lihat semua.")
    add_bullet(doc, "Stok menipis: memberi sinyal produk yang perlu diperiksa di Inventori.")
    add_h2(doc, "Cara memakai dashboard")
    add_number(doc, "Pilih periode yang ingin dianalisis.")
    add_number(doc, "Pilih outlet dari pemilih outlet jika akun memiliki akses multi-outlet.")
    add_number(doc, "Gunakan KPI sebagai ringkasan cepat, lalu buka laporan jika perlu rincian atau export Excel.")

    # POS
    page_break(doc)
    add_section_cover(doc, "4", "POS Kasir dan transaksi", "POS adalah ruang kerja penjualan. Halaman ini dibuat khusus agar kasir dapat fokus melayani pelanggan tanpa menu admin.")
    add_screenshot(doc, "manual-assets-pos.png", "Gambar 6. POS Kasir: status meja, filter meja, pesanan takeaway, dan panel pesanan aktif.")
    add_h2(doc, "Memulai transaksi")
    add_number(doc, "Masuk melalui jalur kasir publik atau klik POS Kasir dari dashboard.")
    add_number(doc, "Pilih meja jika transaksi dine-in. Untuk pesanan tanpa meja, klik Pesanan Baru (Takeaway).")
    add_number(doc, "Pilih kategori lalu produk untuk memasukkan item ke keranjang.")
    add_number(doc, "Atur jumlah item, catatan, varian, atau topping jika tersedia.")
    add_number(doc, "Tinjau total, pajak/service charge, dan diskon yang diizinkan role.")
    add_h2(doc, "Pembayaran dan penyelesaian")
    add_bullet(doc, "Pilih metode pembayaran yang aktif, misalnya Tunai atau QRIS.")
    add_bullet(doc, "Untuk Tunai, masukkan nominal bayar dan pastikan kembalian benar.")
    add_bullet(doc, "Selesaikan transaksi. Sistem menyimpan order dan payment sebagai transaksi selesai.")
    add_bullet(doc, "Gunakan Riwayat Transaksi untuk mencari transaksi yang sudah dibuat atau mencetak ulang struk jika tersedia.")
    add_callout(doc, "Tips kasir", "Selesaikan transaksi setelah pelanggan benar-benar membayar. Untuk koreksi transaksi, gunakan permission void/refund sesuai kebijakan bisnis dan jangan menghapus data transaksi secara langsung.", ORANGE)
    add_h2(doc, "Shift Kasir")
    add_screenshot(doc, "manual-assets-shifts.png", "Gambar 7. Halaman Shift Kasir dengan filter kasir/outlet, kas awal, penjualan, selisih, dan status.")
    add_bullet(doc, "Buka shift dengan menetapkan kas awal sebelum transaksi dimulai.")
    add_bullet(doc, "Pantau total penjualan dan selisih selama shift berjalan.")
    add_bullet(doc, "Tutup shift setelah kas fisik dihitung dan cocokkan dengan nilai yang diharapkan sistem.")

    # Products
    page_break(doc)
    add_section_cover(doc, "5", "Produk dan kategori", "Master produk menjadi sumber data untuk POS, harga, topping/varian, serta laporan produk.")
    add_screenshot(doc, "manual-assets-products.png", "Gambar 8. Halaman Daftar Produk.")
    add_h2(doc, "Produk")
    add_bullet(doc, "Tambah produk: isi nama, kategori, SKU, barcode, gambar, harga modal, harga jual, dan status aktif.")
    add_bullet(doc, "SKU dapat dibuat otomatis sesuai alur aplikasi; barcode dapat diisi jika produk memiliki barcode.")
    add_bullet(doc, "Gunakan status stok untuk membedakan produk yang dikelola stoknya, seperti Aqua, dan produk racikan yang tidak dihitung sebagai satu stok jadi.")
    add_bullet(doc, "Kelola varian/topping untuk produk yang memiliki pilihan ukuran atau tambahan.")
    add_bullet(doc, "Edit atau nonaktifkan produk jika tidak lagi dijual; hindari menghapus produk yang sudah dipakai transaksi jika sistem meminta soft delete.")
    add_h2(doc, "Kategori")
    add_screenshot(doc, "manual-assets-categories.png", "Gambar 9. Halaman Kategori.")
    add_bullet(doc, "Buat kategori seperti Coffee, Non Coffee, Snack, atau kategori operasional lain.")
    add_bullet(doc, "Kategori dipakai untuk filter di POS dan pengelompokan produk di laporan.")
    add_bullet(doc, "Sebelum menghapus kategori, pindahkan produk yang masih menggunakannya.")

    # Inventory
    page_break(doc)
    add_section_cover(doc, "6", "Inventori dan stok", "Inventori mengelola stok per outlet, riwayat mutasi, transfer antar outlet, dan stok opname.")
    add_screenshot(doc, "manual-assets-inventory.png", "Gambar 10. Overview inventori per outlet.")
    add_h2(doc, "Konsep stok Bayaro POS")
    add_body(doc, "Stok melekat pada produk yang ditandai sebagai produk stok. Produk jadi seperti Aqua dapat memiliki stok langsung. Produk racikan seperti Nasi Goreng dapat dijual tanpa stok produk jadi, atau dikembangkan lebih lanjut sebagai bahan baku/recipe jika proses produksi dipakai oleh bisnis.")
    add_h2(doc, "Submodul inventori")
    add_screenshot(doc, "manual-assets-inventory-low-stock.png", "Gambar 11. Daftar stok rendah.")
    add_bullet(doc, "Stok Rendah: memantau item di bawah ambang minimum.")
    add_screenshot(doc, "manual-assets-inventory-adjustments.png", "Gambar 12. Riwayat penyesuaian stok.")
    add_bullet(doc, "Penyesuaian Stok: menambah atau mengurangi stok dengan alasan yang tercatat.")
    add_screenshot(doc, "manual-assets-inventory-transfers.png", "Gambar 13. Transfer stok antar outlet.")
    add_bullet(doc, "Transfer Stok: membuat pengiriman stok dari outlet asal ke outlet tujuan dan memantau statusnya.")
    add_screenshot(doc, "manual-assets-inventory-opname.png", "Gambar 14. Stok opname.")
    add_bullet(doc, "Stok Opname: mencatat hasil hitung fisik dan menyelaraskan stok sistem.")
    add_h2(doc, "Alur stok yang disarankan")
    add_number(doc, "Tetapkan produk mana yang memakai stok.")
    add_number(doc, "Catat penerimaan awal melalui penyesuaian stok dengan alasan yang jelas.")
    add_number(doc, "Periksa Stok Rendah secara berkala.")
    add_number(doc, "Gunakan transfer untuk perpindahan antar outlet dan opname untuk audit fisik.")

    # Customer and promo
    add_section_cover(doc, "7", "Pelanggan dan promo", "CRM membantu menyimpan informasi pelanggan, sementara promo mengatur voucher, diskon, bundle, atau harga berbasis waktu.")
    add_screenshot(doc, "manual-assets-customers.png", "Gambar 15. Halaman Pelanggan.")
    add_h2(doc, "Pelanggan")
    add_bullet(doc, "Tambah pelanggan dengan nama, nomor telepon, email, dan catatan bila diperlukan.")
    add_bullet(doc, "Cari pelanggan saat checkout agar transaksi terhubung ke riwayat kunjungan dan total belanja.")
    add_bullet(doc, "Gunakan detail pelanggan untuk melihat histori transaksi dan catatan layanan.")
    add_screenshot(doc, "manual-assets-promos.png", "Gambar 16. Halaman Promo & Diskon.")
    add_h2(doc, "Promo & Diskon")
    add_bullet(doc, "Buat kode voucher dengan diskon persentase atau nominal tetap.")
    add_bullet(doc, "Atur masa berlaku, minimum transaksi, batas penggunaan, dan status aktif.")
    add_bullet(doc, "Gunakan aturan bundle atau happy hour jika tersedia pada konfigurasi promo.")
    add_bullet(doc, "Uji promo di POS sebelum diumumkan agar syarat dan nilai potongan sesuai.")

    # Reports
    page_break(doc)
    add_section_cover(doc, "8", "Laporan", "Laporan dipakai untuk membaca performa penjualan, produk, kasir, dan inventori dengan filter periode/outlet serta export Excel.")
    add_screenshot(doc, "manual-assets-reports-sales.png", "Gambar 17. Laporan Penjualan dengan KPI, grafik omzet, rincian harian, filter, dan Export Excel.")
    add_h2(doc, "Laporan Penjualan")
    add_bullet(doc, "Melihat total pendapatan, total transaksi, rata-rata transaksi, grafik omzet per hari, dan rincian harian.")
    add_bullet(doc, "Pilih tanggal/periode dan outlet, lalu klik Export Excel untuk mengunduh data.")
    add_screenshot(doc, "manual-assets-reports-products.png", "Gambar 18. Laporan Produk Terlaris.")
    add_h2(doc, "Produk Terlaris")
    add_bullet(doc, "Membandingkan produk berdasarkan jumlah terjual dan pendapatan.")
    add_bullet(doc, "Gunakan laporan ini untuk keputusan menu, bundling, dan pengadaan stok.")
    add_screenshot(doc, "manual-assets-reports-cashier.png", "Gambar 19. Laporan Kasir.")
    add_h2(doc, "Laporan Kasir")
    add_bullet(doc, "Melihat jumlah kasir aktif, transaksi, revenue per kasir, serta detail kinerja kasir.")
    add_bullet(doc, "Bermanfaat untuk evaluasi shift dan rekonsiliasi penjualan.")
    add_screenshot(doc, "manual-assets-reports-inventory.png", "Gambar 20. Laporan Inventori.")
    add_h2(doc, "Laporan Inventori")
    add_bullet(doc, "Menampilkan total produk, stok rendah, stok kritis, serta stok per outlet.")
    add_bullet(doc, "Gunakan bersama Stok Rendah dan Stok Opname untuk tindak lanjut operasional.")

    # Operational
    add_section_cover(doc, "9", "Operasional: karyawan, outlet, dan role", "Modul operasional mengatur siapa yang bekerja, di outlet mana, dan akses apa yang dimiliki.")
    add_screenshot(doc, "manual-assets-employees.png", "Gambar 21. Halaman Karyawan.")
    add_h2(doc, "Karyawan")
    add_bullet(doc, "Tambahkan karyawan dengan nama, email, role, PIN, dan status aktif.")
    add_bullet(doc, "PIN digunakan untuk jalur kasir publik; jaga kerahasiaan PIN per karyawan.")
    add_screenshot(doc, "manual-assets-outlets.png", "Gambar 22. Halaman Outlet.")
    add_h2(doc, "Outlet")
    add_bullet(doc, "Buat dan edit outlet, alamat, kontak, jam operasional, serta status aktif.")
    add_bullet(doc, "Outlet aktif menjadi konteks untuk stok, POS, shift, dan laporan.")
    add_screenshot(doc, "manual-assets-roles.png", "Gambar 23. Halaman Role & Akses.")
    add_h2(doc, "Role & Akses")
    add_bullet(doc, "Buat role khusus jika role bawaan belum cukup.")
    add_bullet(doc, "Pilih permission berdasarkan tanggung jawab, bukan berdasarkan jabatan saja.")
    add_bullet(doc, "Setelah mengubah role, minta pengguna login ulang agar hak akses terbaru terbaca.")

    # Settings
    add_section_cover(doc, "10", "Pengaturan dan akun", "Pengaturan menyimpan konfigurasi bisnis, pajak, pembayaran, struk, printer, serta identitas pengguna.")
    add_screenshot(doc, "manual-assets-settings-business.png", "Gambar 24. Profil Bisnis.")
    add_bullet(doc, "Profil Bisnis: nama bisnis, logo, alamat, telepon, dan informasi branding.")
    add_screenshot(doc, "manual-assets-settings-tax.png", "Gambar 25. Pajak & Service.")
    add_bullet(doc, "Pajak & Service: atur tax rate dan service charge yang dipakai saat checkout.")
    add_screenshot(doc, "manual-assets-settings-receipt.png", "Gambar 26. Template Struk.")
    add_bullet(doc, "Template Struk: header, footer, logo, alamat, telepon, kasir, dan ucapan terima kasih.")
    add_screenshot(doc, "manual-assets-settings-payment.png", "Gambar 27. Metode Bayar.")
    add_bullet(doc, "Metode Bayar: aktifkan/nonaktifkan Tunai, QRIS, EDC, atau gateway yang dipakai.")
    add_screenshot(doc, "manual-assets-settings-printer.png", "Gambar 28. Printer.")
    add_bullet(doc, "Printer: atur koneksi dan preferensi cetak jika perangkat printer tersedia.")
    add_screenshot(doc, "manual-assets-profile.png", "Gambar 29. Profil Saya.")
    add_bullet(doc, "Profil Saya: periksa identitas pengguna yang sedang login.")
    add_screenshot(doc, "manual-assets-settings-account.png", "Gambar 30. Pengaturan Akun.")
    add_bullet(doc, "Pengaturan Akun: kelola informasi akun dan pengaturan personal yang tersedia.")

    # Daily flow and troubleshooting
    page_break(doc)
    add_section_cover(doc, "11", "Alur kerja harian dan troubleshooting", "Bagian ini merangkum urutan kerja yang aman dan cara membaca masalah umum.")
    add_h2(doc, "Checklist pembukaan hari")
    for text in [
        "Pastikan outlet yang dipilih benar.",
        "Pastikan shift kasir dibuka dan kas awal dicatat.",
        "Periksa produk aktif, harga jual, metode pembayaran, dan stok rendah.",
        "Pastikan printer atau perangkat pembayaran siap jika digunakan.",
    ]:
        add_bullet(doc, text)
    add_h2(doc, "Checklist saat operasional")
    for text in [
        "Gunakan produk yang benar dan pilih varian/topping sesuai pesanan.",
        "Pastikan pembayaran selesai sebelum transaksi ditutup.",
        "Gunakan void/refund hanya sesuai otorisasi role.",
        "Catat koreksi stok dengan alasan yang jelas.",
    ]:
        add_bullet(doc, text)
    add_h2(doc, "Checklist penutupan hari")
    for text in [
        "Buka Shift Kasir dan lakukan rekonsiliasi uang tunai.",
        "Bandingkan ringkasan penjualan dengan Laporan Penjualan.",
        "Periksa Stok Rendah dan jadwalkan pengadaan/transfer.",
        "Tutup shift dan simpan catatan selisih jika ada.",
    ]:
        add_bullet(doc, text)
    add_h2(doc, "Masalah umum")
    add_table(doc, ["Gejala", "Pemeriksaan", "Tindakan"], [
        ["Tidak bisa login", "Email/password dan status akun", "Gunakan akun yang benar, reset password jika tersedia, dan cek environment/auth secret."],
        ["Menu tidak terlihat", "Role dan permission", "Periksa Role & Akses lalu login ulang."],
        ["POS meminta PIN", "Outlet dan PIN karyawan", "Pastikan kode toko, outlet, dan PIN sesuai."],
        ["Data stok tidak sesuai", "Outlet aktif dan mutasi terakhir", "Periksa Penyesuaian, Transfer, dan Stok Opname."],
        ["Halaman error setelah perubahan", "Terminal server dan build", "Restart server, bersihkan cache `.next` bila perlu, lalu jalankan `npm run dev` kembali."],
    ], [1.55, 2.0, 2.8])
    add_callout(doc, "Server lokal", "Untuk menjalankan aplikasi: npm install → npx prisma migrate deploy → npx prisma db seed → npm run dev. Aplikasi tersedia di http://localhost:3000.", PALE_BLUE)
    add_body(doc, "Manual book ini menggunakan data demo dan screenshot aktual dari aplikasi. Label atau jumlah data dapat berubah mengikuti isi database dan konfigurasi outlet.", space_after=0, color=MUTED, size=9)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
